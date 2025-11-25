#!/usr/bin/env python3
"""
Advanced Export System

Comprehensive export functionality supporting PDF, HTML, DOCX, and other formats
with customizable templates and professional styling.
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
import tempfile
import base64

# Try to import export dependencies with fallbacks
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("ReportLab not available - PDF export disabled")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available - DOCX export disabled")

try:
    from jinja2 import Template, Environment, FileSystemLoader
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    print("Jinja2 not available - HTML export disabled")

@dataclass
class ExportOptions:
    """Configuration options for exports."""
    format: str  # 'pdf', 'html', 'docx', 'json', 'markdown'
    template: Optional[str] = None  # Custom template name
    include_charts: bool = True
    include_code_samples: bool = True
    include_recommendations: bool = True
    company_name: Optional[str] = None
    project_name: Optional[str] = None
    author: Optional[str] = None
    custom_css: Optional[str] = None
    theme: str = 'professional'  # 'professional', 'modern', 'minimal'

class AdvancedExporter:
    """
    Advanced export system supporting multiple formats with professional styling.
    
    Features:
    - PDF reports with charts and tables
    - HTML reports with interactive elements
    - DOCX documents with styling
    - JSON structured data
    - Markdown documentation
    - Custom templates and themes
    - Batch export capabilities
    """
    
    def __init__(self, template_dir: Optional[Path] = None):
        """Initialize exporter with template directory."""
        self.template_dir = template_dir or Path(__file__).parent / 'templates'
        self.template_dir.mkdir(exist_ok=True)
        
        # Initialize template environment
        if HTML_AVAILABLE:
            self.jinja_env = Environment(
                loader=FileSystemLoader(str(self.template_dir)),
                autoescape=True
            )
        
        # Create default templates if they don't exist
        self._ensure_default_templates()
    
    def export_analysis(self, 
                       analysis_result: Dict[str, Any], 
                       options: ExportOptions,
                       output_path: Optional[Path] = None) -> Path:
        """
        Export analysis result in specified format.
        
        Args:
            analysis_result: Analysis result dictionary
            options: Export configuration options
            output_path: Optional output file path
            
        Returns:
            Path to exported file
        """
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"analysis_report_{timestamp}.{options.format}"
            output_path = Path.cwd() / filename
        
        # Prepare data for export
        export_data = self._prepare_export_data(analysis_result, options)
        
        # Export based on format
        if options.format.lower() == 'pdf':
            return self._export_pdf(export_data, options, output_path)
        elif options.format.lower() == 'html':
            return self._export_html(export_data, options, output_path)
        elif options.format.lower() == 'docx':
            return self._export_docx(export_data, options, output_path)
        elif options.format.lower() == 'json':
            return self._export_json(export_data, options, output_path)
        elif options.format.lower() == 'markdown':
            return self._export_markdown(export_data, options, output_path)
        else:
            raise ValueError(f"Unsupported export format: {options.format}")
    
    def _prepare_export_data(self, analysis_result: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Prepare and enrich data for export."""
        # Calculate summary statistics
        all_issues = []
        if 'critical_issues' in analysis_result:
            all_issues.extend(analysis_result['critical_issues'])
        if 'high_priority_issues' in analysis_result:
            all_issues.extend(analysis_result['high_priority_issues'])
        if 'medium_priority_issues' in analysis_result:
            all_issues.extend(analysis_result['medium_priority_issues'])
        if 'low_priority_issues' in analysis_result:
            all_issues.extend(analysis_result['low_priority_issues'])
        
        # Group issues by type
        issues_by_type = {}
        issues_by_severity = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0}
        
        for issue in all_issues:
            issue_type = issue.get('type', 'unknown')
            if issue_type not in issues_by_type:
                issues_by_type[issue_type] = 0
            issues_by_type[issue_type] += 1
            
            severity = issue.get('severity', 'unknown')
            if severity in issues_by_severity:
                issues_by_severity[severity] += 1
        
        # Prepare enriched data
        export_data = {
            **analysis_result,
            'export_metadata': {
                'exported_at': datetime.now().isoformat(),
                'export_format': options.format,
                'theme': options.theme,
                'company_name': options.company_name or 'Prompt Engineer',
                'project_name': options.project_name or Path(analysis_result.get('project_path', '')).name,
                'author': options.author or 'Prompt Engineer Analysis Tool'
            },
            'summary_stats': {
                'total_issues': len(all_issues),
                'issues_by_severity': issues_by_severity,
                'issues_by_type': dict(sorted(issues_by_type.items(), key=lambda x: x[1], reverse=True)),
                'health_score': analysis_result.get('health_score', 0)
            }
        }
        
        return export_data
    
    def _export_pdf(self, data: Dict[str, Any], options: ExportOptions, output_path: Path) -> Path:
        """Export to PDF format."""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF export not available - install reportlab")
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=18
        )
        
        # Build story (content)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        # Title page
        story.append(Paragraph(f"Project Analysis Report", title_style))
        story.append(Spacer(1, 20))
        
        metadata = data['export_metadata']
        story.append(Paragraph(f"<b>Project:</b> {metadata['project_name']}", styles['Normal']))
        story.append(Paragraph(f"<b>Analyzed:</b> {metadata['exported_at'][:19]}", styles['Normal']))
        story.append(Paragraph(f"<b>Health Score:</b> {data.get('health_score', 'N/A')}/100", styles['Normal']))
        story.append(PageBreak())
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        summary_stats = data['summary_stats']
        
        summary_data = [
            ['Metric', 'Value'],
            ['Overall Health Score', f"{data.get('health_score', 'N/A')}/100"],
            ['Total Issues Found', str(summary_stats['total_issues'])],
            ['Critical Issues', str(summary_stats['issues_by_severity']['critical'])],
            ['High Priority Issues', str(summary_stats['issues_by_severity']['high'])],
            ['Medium Priority Issues', str(summary_stats['issues_by_severity']['medium'])],
            ['Low Priority Issues', str(summary_stats['issues_by_severity']['low'])]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # Issues by category
        if options.include_recommendations and summary_stats['issues_by_type']:
            story.append(Paragraph("Issues by Category", heading_style))
            
            category_data = [['Issue Type', 'Count']]
            for issue_type, count in list(summary_stats['issues_by_type'].items())[:10]:  # Top 10
                category_data.append([issue_type.replace('_', ' ').title(), str(count)])
            
            category_table = Table(category_data, colWidths=[3*inch, 2*inch])
            category_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(category_table)
            story.append(PageBreak())
        
        # Detailed issues
        for severity in ['critical_issues', 'high_priority_issues']:
            if severity in data and data[severity]:
                severity_name = severity.replace('_', ' ').title()
                story.append(Paragraph(f"{severity_name}", heading_style))
                
                for i, issue in enumerate(data[severity][:5]):  # Top 5 per severity
                    story.append(Paragraph(f"<b>{i+1}. {issue.get('title', 'Untitled Issue')}</b>", styles['Normal']))
                    story.append(Paragraph(f"<i>{issue.get('description', 'No description')}</i>", styles['Normal']))
                    if issue.get('suggested_action'):
                        story.append(Paragraph(f"<b>Action:</b> {issue['suggested_action']}", styles['Normal']))
                    story.append(Spacer(1, 12))
                
                story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        return output_path
    
    def _export_html(self, data: Dict[str, Any], options: ExportOptions, output_path: Path) -> Path:
        """Export to HTML format."""
        if not HTML_AVAILABLE:
            # Fallback basic HTML
            html_content = self._generate_basic_html(data, options)
        else:
            # Use Jinja2 template
            template_name = options.template or 'default_report.html'
            try:
                template = self.jinja_env.get_template(template_name)
                html_content = template.render(data=data, options=options)
            except:
                # Fallback to basic HTML if template fails
                html_content = self._generate_basic_html(data, options)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def _generate_basic_html(self, data: Dict[str, Any], options: ExportOptions) -> str:
        """Generate basic HTML report without Jinja2."""
        metadata = data['export_metadata']
        summary_stats = data['summary_stats']
        
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Project Analysis Report - {metadata['project_name']}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #2c3e50; text-align: center; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; border-left: 4px solid #3498db; padding-left: 15px; }}
        .summary {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; margin: 20px 0; }}
        .metric {{ background: #ecf0f1; padding: 15px; border-radius: 8px; text-align: center; }}
        .metric h3 {{ margin: 0; color: #2c3e50; }}
        .metric p {{ font-size: 24px; font-weight: bold; margin: 5px 0; }}
        .health-score {{ background: linear-gradient(135deg, #27ae60, #2ecc71); color: white; }}
        .critical {{ background: linear-gradient(135deg, #e74c3c, #c0392b); color: white; }}
        .high {{ background: linear-gradient(135deg, #f39c12, #e67e22); color: white; }}
        .issues {{ margin: 20px 0; }}
        .issue {{ background: #f8f9fa; margin: 10px 0; padding: 15px; border-left: 4px solid #3498db; border-radius: 4px; }}
        .issue h4 {{ margin: 0 0 10px 0; color: #2c3e50; }}
        .issue p {{ margin: 5px 0; color: #7f8c8d; }}
        .footer {{ margin-top: 40px; text-align: center; color: #95a5a6; font-size: 12px; }}
        @media print {{ body {{ background: white; }} .container {{ box-shadow: none; }} }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Project Analysis Report</h1>
        
        <div class="summary">
            <div class="metric health-score">
                <h3>Health Score</h3>
                <p>{data.get('health_score', 'N/A')}/100</p>
            </div>
            <div class="metric">
                <h3>Total Issues</h3>
                <p>{summary_stats['total_issues']}</p>
            </div>
            <div class="metric critical">
                <h3>Critical</h3>
                <p>{summary_stats['issues_by_severity']['critical']}</p>
            </div>
            <div class="metric high">
                <h3>High Priority</h3>
                <p>{summary_stats['issues_by_severity']['high']}</p>
            </div>
        </div>
        
        <h2>Project Information</h2>
        <p><strong>Project:</strong> {metadata['project_name']}</p>
        <p><strong>Analyzed:</strong> {metadata['exported_at'][:19]}</p>
        <p><strong>Technology Stack:</strong> {', '.join(data.get('tech_stack', []))}</p>
        """
        
        # Add critical and high priority issues
        for severity in ['critical_issues', 'high_priority_issues']:
            if severity in data and data[severity]:
                severity_name = severity.replace('_', ' ').title()
                html += f"<h2>{severity_name}</h2><div class='issues'>"
                
                for issue in data[severity][:10]:  # Top 10
                    html += f"""
                    <div class="issue">
                        <h4>{issue.get('title', 'Untitled Issue')}</h4>
                        <p><strong>Description:</strong> {issue.get('description', 'No description')}</p>
                    """
                    if issue.get('suggested_action'):
                        html += f"<p><strong>Action:</strong> {issue['suggested_action']}</p>"
                    html += "</div>"
                
                html += "</div>"
        
        html += f"""
        <div class="footer">
            <p>Generated by {metadata['company_name']} • {metadata['exported_at'][:19]}</p>
        </div>
    </div>
</body>
</html>
        """
        
        return html
    
    def _export_docx(self, data: Dict[str, Any], options: ExportOptions, output_path: Path) -> Path:
        """Export to DOCX format."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX export not available - install python-docx")
        
        doc = Document()
        
        # Add title
        metadata = data['export_metadata']
        title = doc.add_heading('Project Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add project info
        doc.add_paragraph(f"Project: {metadata['project_name']}")
        doc.add_paragraph(f"Analyzed: {metadata['exported_at'][:19]}")
        doc.add_paragraph(f"Health Score: {data.get('health_score', 'N/A')}/100")
        
        # Add summary
        doc.add_heading('Executive Summary', level=1)
        summary_stats = data['summary_stats']
        
        p = doc.add_paragraph()
        p.add_run(f"Total Issues: {summary_stats['total_issues']}").bold = True
        p.add_run(f"\nCritical: {summary_stats['issues_by_severity']['critical']}")
        p.add_run(f"\nHigh Priority: {summary_stats['issues_by_severity']['high']}")
        p.add_run(f"\nMedium Priority: {summary_stats['issues_by_severity']['medium']}")
        p.add_run(f"\nLow Priority: {summary_stats['issues_by_severity']['low']}")
        
        # Add issues
        for severity in ['critical_issues', 'high_priority_issues']:
            if severity in data and data[severity]:
                severity_name = severity.replace('_', ' ').title()
                doc.add_heading(severity_name, level=1)
                
                for i, issue in enumerate(data[severity][:5]):  # Top 5
                    doc.add_heading(f"{i+1}. {issue.get('title', 'Untitled Issue')}", level=2)
                    doc.add_paragraph(issue.get('description', 'No description'))
                    
                    if issue.get('suggested_action'):
                        p = doc.add_paragraph()
                        p.add_run('Recommended Action: ').bold = True
                        p.add_run(issue['suggested_action'])
        
        doc.save(output_path)
        return output_path
    
    def _export_json(self, data: Dict[str, Any], options: ExportOptions, output_path: Path) -> Path:
        """Export to JSON format."""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, default=str, ensure_ascii=False)
        
        return output_path
    
    def _export_markdown(self, data: Dict[str, Any], options: ExportOptions, output_path: Path) -> Path:
        """Export to Markdown format."""
        metadata = data['export_metadata']
        summary_stats = data['summary_stats']
        
        md_content = f"""# Project Analysis Report

**Project:** {metadata['project_name']}  
**Analyzed:** {metadata['exported_at'][:19]}  
**Health Score:** {data.get('health_score', 'N/A')}/100

## Executive Summary

| Metric | Value |
|--------|--------|
| Total Issues | {summary_stats['total_issues']} |
| Critical Issues | {summary_stats['issues_by_severity']['critical']} |
| High Priority Issues | {summary_stats['issues_by_severity']['high']} |
| Medium Priority Issues | {summary_stats['issues_by_severity']['medium']} |
| Low Priority Issues | {summary_stats['issues_by_severity']['low']} |

## Technology Stack

{', '.join(data.get('tech_stack', []))}

"""
        
        # Add issues
        for severity in ['critical_issues', 'high_priority_issues']:
            if severity in data and data[severity]:
                severity_name = severity.replace('_', ' ').title()
                md_content += f"## {severity_name}\n\n"
                
                for i, issue in enumerate(data[severity][:10]):
                    md_content += f"### {i+1}. {issue.get('title', 'Untitled Issue')}\n\n"
                    md_content += f"{issue.get('description', 'No description')}\n\n"
                    
                    if issue.get('suggested_action'):
                        md_content += f"**Recommended Action:** {issue['suggested_action']}\n\n"
                    
                    md_content += "---\n\n"
        
        md_content += f"\n*Generated by {metadata['company_name']} • {metadata['exported_at'][:19]}*\n"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return output_path
    
    def _ensure_default_templates(self):
        """Create default templates if they don't exist."""
        default_template_path = self.template_dir / 'default_report.html'
        
        if HTML_AVAILABLE and not default_template_path.exists():
            default_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Analysis Report - {{ data.export_metadata.project_name }}</title>
    <style>
        /* Professional styling */
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 0; background: #f8f9fa; }
        .container { max-width: 1200px; margin: 20px auto; background: white; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 20px rgba(0,0,0,0.1); }
        .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 40px; text-align: center; }
        .content { padding: 40px; }
        h1 { margin: 0; font-size: 2.5em; font-weight: 300; }
        h2 { color: #4a5568; border-bottom: 2px solid #e2e8f0; padding-bottom: 10px; }
        .metrics { display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin: 30px 0; }
        .metric-card { background: #f7fafc; padding: 20px; border-radius: 8px; border-left: 4px solid #667eea; }
        .metric-value { font-size: 2em; font-weight: bold; color: #2d3748; }
        .issue { background: #fff5f5; border: 1px solid #feb2b2; border-radius: 6px; padding: 15px; margin: 10px 0; }
        .issue.critical { border-color: #fc8181; background: #fed7d7; }
        .issue.high { border-color: #f6ad55; background: #feebc8; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{{ data.export_metadata.project_name }}</h1>
            <p>Analysis Report • {{ data.export_metadata.exported_at[:19] }}</p>
        </div>
        
        <div class="content">
            <div class="metrics">
                <div class="metric-card">
                    <h3>Health Score</h3>
                    <div class="metric-value">{{ data.health_score or 'N/A' }}/100</div>
                </div>
                <div class="metric-card">
                    <h3>Total Issues</h3>
                    <div class="metric-value">{{ data.summary_stats.total_issues }}</div>
                </div>
                <div class="metric-card">
                    <h3>Critical Issues</h3>
                    <div class="metric-value">{{ data.summary_stats.issues_by_severity.critical }}</div>
                </div>
            </div>
            
            {% if data.critical_issues %}
            <h2>Critical Issues</h2>
            {% for issue in data.critical_issues[:5] %}
            <div class="issue critical">
                <h4>{{ issue.title or 'Untitled Issue' }}</h4>
                <p>{{ issue.description or 'No description' }}</p>
                {% if issue.suggested_action %}
                <p><strong>Action:</strong> {{ issue.suggested_action }}</p>
                {% endif %}
            </div>
            {% endfor %}
            {% endif %}
        </div>
    </div>
</body>
</html>"""
            
            with open(default_template_path, 'w', encoding='utf-8') as f:
                f.write(default_template)
    
    def batch_export(self, 
                    analysis_results: List[Dict[str, Any]], 
                    formats: List[str],
                    output_dir: Path,
                    base_options: ExportOptions) -> List[Path]:
        """Export multiple analysis results to multiple formats."""
        exported_files = []
        
        output_dir.mkdir(exist_ok=True)
        
        for i, result in enumerate(analysis_results):
            project_name = Path(result.get('project_path', f'project_{i}')).name
            
            for format_type in formats:
                options = ExportOptions(
                    format=format_type,
                    template=base_options.template,
                    include_charts=base_options.include_charts,
                    include_code_samples=base_options.include_code_samples,
                    include_recommendations=base_options.include_recommendations,
                    company_name=base_options.company_name,
                    project_name=project_name,
                    author=base_options.author,
                    theme=base_options.theme
                )
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{project_name}_{timestamp}.{format_type}"
                output_path = output_dir / filename
                
                try:
                    exported_file = self.export_analysis(result, options, output_path)
                    exported_files.append(exported_file)
                except Exception as e:
                    print(f"Failed to export {filename}: {e}")
        
        return exported_files